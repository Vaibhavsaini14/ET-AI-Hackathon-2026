import pymupdf
import os
from dataclasses import dataclass, field
from loguru import logger

@dataclass
class ParsedDocument:
    doc_id: str
    file_path: str
    doc_type: str
    title: str
    raw_text: str
    pages: list = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

class DocumentParser:

    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        extension = os.path.splitext(file_path)[1].lower()
        title = os.path.splitext(os.path.basename(file_path))[0]

        if extension == ".pdf":
            return self._parse_pdf(file_path, doc_id, title)
        elif extension == ".txt":
            return self._parse_txt(file_path, doc_id, title)
        elif extension == ".docx":
            return self._parse_docx(file_path, doc_id, title)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def _parse_pdf(self, file_path: str, doc_id: str, title: str) -> ParsedDocument:
        logger.info(f"Parsing PDF: {title}")
        doc = pymupdf.open(file_path)
        pages = []
        full_text = ""

        for page_num, page in enumerate(doc):
            text = page.get_text("text")

            if len(text.strip()) < 30:
                logger.warning(f"Page {page_num + 1} has very little text, may be scanned")

            pages.append({
                "page_num": page_num + 1,
                "text": text,
                "char_count": len(text)
            })
            full_text += f"\n\n--- PAGE {page_num + 1} ---\n{text}"

        doc.close()
        logger.info(f"Parsed {len(pages)} pages from {title}")

        return ParsedDocument(
            doc_id=doc_id,
            file_path=file_path,
            doc_type="unknown",
            title=title,
            raw_text=full_text,
            pages=pages,
            metadata={
                "page_count": len(pages),
                "file_size_kb": os.path.getsize(file_path) // 1024,
                "extension": ".pdf"
            }
        )

    def _parse_txt(self, file_path: str, doc_id: str, title: str) -> ParsedDocument:
        logger.info(f"Parsing TXT: {title}")
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        return ParsedDocument(
            doc_id=doc_id,
            file_path=file_path,
            doc_type="unknown",
            title=title,
            raw_text=text,
            pages=[{"page_num": 1, "text": text, "char_count": len(text)}],
            metadata={"page_count": 1, "extension": ".txt"}
        )

    def _parse_docx(self, file_path: str, doc_id: str, title: str) -> ParsedDocument:
        logger.info(f"Parsing DOCX: {title}")
        from docx import Document
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        return ParsedDocument(
            doc_id=doc_id,
            file_path=file_path,
            doc_type="unknown",
            title=title,
            raw_text=text,
            pages=[{"page_num": 1, "text": text, "char_count": len(text)}],
            metadata={"page_count": 1, "extension": ".docx"}
        )